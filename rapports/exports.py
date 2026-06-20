"""
Utilitaires d'export : PDF (ReportLab), Excel (openpyxl), Word (python-docx).
"""
import io
from datetime import date

from django.http import HttpResponse


# ─── EXCEL ──────────────────────────────────────────────────────────────────

def export_excel(titre, entetes, lignes, nom_fichier='rapport.xlsx'):
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = titre[:30]

    # Titre
    ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=len(entetes))
    ws['A1'] = titre
    ws['A1'].font = Font(bold=True, size=14)
    ws['A1'].alignment = Alignment(horizontal='center')

    # En-têtes
    bleu = PatternFill(fill_type='solid', fgColor='1F4E79')
    for col, entete in enumerate(entetes, 1):
        cell = ws.cell(row=2, column=col, value=entete)
        cell.font = Font(bold=True, color='FFFFFF')
        cell.fill = bleu
        cell.alignment = Alignment(horizontal='center')

    # Données
    for row_idx, ligne in enumerate(lignes, 3):
        for col_idx, val in enumerate(ligne, 1):
            ws.cell(row=row_idx, column=col_idx, value=val)

    # Auto-width
    for col in ws.columns:
        max_len = max((len(str(c.value or '')) for c in col), default=10)
        ws.column_dimensions[col[0].column_letter].width = min(max_len + 4, 50)

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    response = HttpResponse(buf, content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = f'attachment; filename="{nom_fichier}"'
    return response


# ─── PDF ────────────────────────────────────────────────────────────────────

def export_pdf(titre, entetes, lignes, nom_fichier='rapport.pdf'):
    from reportlab.lib.pagesizes import A4, landscape
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib.units import cm

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=landscape(A4), topMargin=1.5*cm, bottomMargin=1.5*cm)
    styles = getSampleStyleSheet()
    elements = []

    elements.append(Paragraph(titre, styles['Title']))
    elements.append(Paragraph(f'Généré le {date.today().strftime("%d/%m/%Y")}', styles['Normal']))
    elements.append(Spacer(1, 0.5*cm))

    data = [entetes] + [list(l) for l in lignes]
    t = Table(data, repeatRows=1)
    t.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1F4E79')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('FONTSIZE', (0, 0), (-1, -1), 8),
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#EBF3FB')]),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
        ('TOPPADDING', (0, 0), (-1, -1), 4),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
    ]))
    elements.append(t)
    doc.build(elements)

    buf.seek(0)
    response = HttpResponse(buf, content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{nom_fichier}"'
    return response


# ─── WORD ───────────────────────────────────────────────────────────────────

def export_word(titre, entetes, lignes, nom_fichier='rapport.docx'):
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn
    import docx.oxml

    doc = Document()
    doc.add_heading(titre, 0)
    doc.add_paragraph(f'Généré le {date.today().strftime("%d/%m/%Y")}')

    table = doc.add_table(rows=1 + len(lignes), cols=len(entetes))
    table.style = 'Table Grid'

    # En-têtes
    hdr_row = table.rows[0]
    for i, entete in enumerate(entetes):
        cell = hdr_row.cells[i]
        cell.text = entete
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        # Fond bleu
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = docx.oxml.OxmlElement('w:shd')
        shd.set(qn('w:fill'), '1F4E79')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:val'), 'clear')
        tcPr.append(shd)

    # Données
    for row_idx, ligne in enumerate(lignes, 1):
        row = table.rows[row_idx]
        for col_idx, val in enumerate(ligne):
            row.cells[col_idx].text = str(val or '')

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    response = HttpResponse(
        buf,
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = f'attachment; filename="{nom_fichier}"'
    return response
